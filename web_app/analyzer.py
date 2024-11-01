import re
import os
import copy
import math
import jinja2
from flask import render_template
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.style import WD_STYLE_TYPE

from uniparser_albanian import AlbanianAnalyzer
from uniparser_beserman_lat import BesermanLatAnalyzer
from uniparser_buryat import BuryatAnalyzer
from uniparser_eastern_armenian import EasternArmenianAnalyzer
from uniparser_erzya import ErzyaAnalyzer
from uniparser_komi_zyrian import KomiZyrianAnalyzer
from uniparser_meadow_mari import MeadowMariAnalyzer
from uniparser_moksha import MokshaAnalyzer
from uniparser_ossetic import OsseticAnalyzer
from uniparser_turoyo import TuroyoAnalyzer
from uniparser_udmurt import UdmurtAnalyzer
from uniparser_urmi import UrmiAnalyzer

from .translit_armenian import armenian_translit_meillet
from .translit_beserman import beserman_translit_cyrillic, beserman_translit_upa, beserman_translit_ipa
from .translit_erzya import erzya_translit_upa
from .translit_udmurt import udmurt_translit_upa


class Analyzer:
    rxWords = re.compile('\\w+|\\w[\\w\'-]+\\w|[^\\w]+')
    rxWord = re.compile('\\w+|\\w[\\w\'-]+\\w')
    rxSpace = re.compile('^[ \r\n\t]+$')
    rxBadChars = re.compile('[<>&]')

    def __init__(self):
        self.langs = {
            # 'albanian': {
            #     'name': 'Albanian',
            #     'analyzer': AlbanianAnalyzer()
            # },
            'beserman': {
                'name': 'Beserman (Latin-based)',
                'analyzer': BesermanLatAnalyzer(),
                'translit': {
                    'UPA': beserman_translit_upa,
                    'IPA': beserman_translit_ipa,
                    'Cyrillic': beserman_translit_cyrillic
                }
            }
            # 'buryat': {
            #     'name': 'Buryat',
            #     'analyzer': BuryatAnalyzer()
            # },
            # 'eastern_armenian': {
            #     'name': 'Eastern Armenian',
            #     'analyzer': EasternArmenianAnalyzer(),
            #     'translit': {
            #         'Quasi-Meillet': armenian_translit_meillet
            #     }
            # },
            # 'erzya': {
            #     'name': 'Erzya',
            #     'analyzer': ErzyaAnalyzer(),
            #     'translit': {
            #         'UPA': erzya_translit_upa
            #     }
            # },
            # 'komi_zyrian': {
            #     'name': 'Komi Zyrian',
            #     'analyzer': KomiZyrianAnalyzer()
            # },
            # 'meadow_mari': {
            #     'name': 'Meadow Mari',
            #     'analyzer': MeadowMariAnalyzer()
            # },
            # 'moksha': {
            #     'name': 'Moksha',
            #     'analyzer': MokshaAnalyzer()
            # },
            # 'ossetic': {
            #     'name': 'Ossetic (Iron)',
            #     'analyzer': OsseticAnalyzer()
            # },
            # 'turoyo': {
            #     'name': 'Ṭuroyo',
            #     'analyzer': TuroyoAnalyzer()
            # },
            # 'udmurt': {
            #     'name': 'Udmurt',
            #     'analyzer': UdmurtAnalyzer(),
            #     'translit': {
            #         'UPA': udmurt_translit_upa
            #     }
            # },
            # 'urmi': {
            #     'name': 'Christian Urmi (Assyrian Neo-Aramaic), Latin-based',
            #     'analyzer': UrmiAnalyzer()
            # }
        }
        self.disamb_langs = ['albanian', 'udmurt', 'beserman', 'eastern_armenian']

    def analyze(self, lang, sentence):
        if lang not in self.langs:
            return ''
        sentence = self.rxBadChars.sub('', sentence)[:2048]
        tokens = [t.strip() for t in self.rxWords.findall(sentence.strip())
                  if self.rxSpace.search(t) is None]
        result = []
        if lang in self.disamb_langs:
            result = self.langs[lang]['analyzer'].analyze_words(tokens, disambiguate=True, format='json')
        else:
            result = self.langs[lang]['analyzer'].analyze_words(tokens, format='json')
        result = {'default': result}
        if 'translit' in self.langs[lang]:
            for translit, f in self.langs[lang]['translit'].items():
                resultTranslit = []
                for w in result['default']:
                    wTranslit = copy.deepcopy(w)
                    for ana in wTranslit:
                        ana['wf'] = f(ana['wf'])
                        if 'lemma' in ana:
                            ana['lemma'] = f(ana['lemma'])
                        if 'wfGlossed' in ana:
                            ana['wfGlossed'] = f(ana['wfGlossed'])
                    resultTranslit.append(wTranslit)
                result[translit] = resultTranslit
        return result


class PaperParser:
    rxPuncR = re.compile('^[.,?!:;)"/\\-\\]”]+$')
    rxPuncL = re.compile('^[#*(\\[“]+$')
    rxSingleQuoteL = re.compile('(?<=[ \t\r\n(\\[])\'', flags=re.DOTALL)
    rxSingleQuoteR = re.compile('(?<=[\\w.,?!:;)\\]])\'', flags=re.DOTALL)
    rxDoubleQuoteL = re.compile('(?<=[ \t\r\n(\\[])"', flags=re.DOTALL)
    rxDoubleQuoteR = re.compile('(?<=[\\w.,?!:;)\\]])"', flags=re.DOTALL)
    rxFigureDash = re.compile('(?<=[0-9]) ?- ?(?=[0-9])')
    rxEmDash = re.compile('(?<=[^0-9 ]) - (?=[^0-9 ])')
    rxExamples = re.compile('(?<=\n) *\\((x[x0-9]+?)\\)[ \t]*([^\r\n]+?)[ \t]*\n'
                            '(?: *([^ \r\n][^\r\n]*?) *\n)?|(?<=\n)([^\n]*\n)',
                            flags=re.DOTALL)
    rxStemGloss = re.compile('[ ,;:()]+')
    rxWordLang = {
        'beserman': re.compile('(?<= )-[\\w(́)]*[əɤʼčšžǯɨ́ʉ̯ʌɘʲ͡ɕʂʐʑˌа-яёӵӝӟӥӧʙ̥ʔ̩̥ː][\\ẃ()-]*|'
                               '[-\\ẃ]*[əɤʼčšžǯɨ́ʉ̯ʌɘʲ͡ɕʂʐʑˌа-яёӵӝӟӥӧʙ̥ʔ̩̥ː][-\\ẃ]*|'
                               '(?<= )-[\\w()-]+-(?= )|(?<= )-[\\w()-]+\\b|'
                               '\\b(ta|[mt]on|ben|uk|mare?|(ta|so)os(len)?|nu|(ta|so)je|'
                               'das|og|og-og[^ \r\n]*|kud|kudi[^ \r\n]*|perv[oi]j[^ \r\n]*|'
                               'pe|val|palaz|u[gmzd]|na|tak|odig[^ \r\n]*|se?re|ma|ik|mh|vot|tare|'
                               'ke|ja|bere|pun[eoi]?[mdz]?|gine|(so|ta)iz[^ \r\n]*|gord|marke|'
                               'e[jzmd]|(ta|so)len|(ta|so)(in|len|tek)|tros|bur|luoz|naverno|'
                               'pi|dore|vaj[eo]?|med|da|wa|olo|abi(len)?|jun|\\w+jos(len)?|\\wjez(len)?|'
                               'korka[^ \r\n]*|aslam|poti[zdm]?|kule|lue|murt[^ \r\n]*)\\b',
                               flags=re.DOTALL)
    }
    rxEnlitics = {
        'beserman': re.compile('^([gk][iʌ]ne|uk|ik|nʲi|vedʲ|ʐe|to|no|na|ʂatʲ|ke|pe|a)$', flags=re.I)
    }
    rxGlosses = {
        'beserman': re.compile('\\b(IDEO|REP|AUTOREP|ENIM|(?<![‘\'])ID(?!=\\.)|'
                               'IAM|Q|IMP(?:\\.MTG)?|PROH|HESIT|COMPL|EXIST|'
                               'PRS|PST(?:\\.EVID(?:\\.NEG(?:\\.[123]+)?(?:\\.?(?:SG|PL))?)?)?|'
                               'FUT|(?:ACC\\.)?[123](?:SG|PL)(?:\\.POSS)?|NOT\\.EXIST|'
                               'INDEF|ITER|DETR|CAUS|NOM|GEN2?|ACC(?:\\.PL)?|DAT|INS|(?<![ ‘\'])CAR|ADV|'
                               'LOC|LAT|EL|PROL|EGR|TERM|APP|RCS|DMS|OPT(?:\\.[123])?|'
                               'NEG(?:\\.(?:FUT|PRS|PST))?(?:\\.[123]+)?(?:\\.?(?:SG|PL))?|'
                               'CNG(?:\\.(?:FUT|PRS|PST))?(?:\\.[123]+)?(?:\\.?(?:SG|PL))?|'
                               'COND|COMP|PROP|ATTR|MULT|INF(?:\\.CESS)?|RES|DEB|'
                               'NMLZ|PTCP(?:\\.(?:ACT|NEG|PST|HAB|DEB))?(?:\\.NEG)?|'
                               'ORD|ADVLOC|ADVTEMP|EXHST|DELIM|APPRNUM|RUS|PPF|'
                               'EXCL|INCL|(?<![ ‘\'(])ADD|CONTR|'
                               'CVB(?:\\.(?:NEG|SIM[1-5]?|LIM|REAS\\.NEG))?|PL(?:\\.ADJ)?|SG)\\b',
                               flags=re.DOTALL|re.I)
    }
    rxGlossesNonGlosses = re.compile('([^$]+)')

    @staticmethod
    def clean_punc(text):
        """
        Replace quotation marks and whatnot.
        """
        text = PaperParser.rxSingleQuoteL.sub('‘', text)
        text = PaperParser.rxSingleQuoteR.sub('’', text)
        text = PaperParser.rxDoubleQuoteL.sub('“', text)
        text = PaperParser.rxDoubleQuoteR.sub('”', text)
        text = PaperParser.rxFigureDash.sub('–', text)
        text = PaperParser.rxEmDash.sub('—', text)
        return text

    @staticmethod
    def set_cell_margins(table, left=0, right=0):
        tc = table._element
        tblPr = tc.tblPr
        tblCellMar = OxmlElement('w:tblCellMar')
        kwargs = {"left": left, "right": right}
        for m in ["left", "right"]:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)

        tblPr.append(tblCellMar)

    @staticmethod
    def p_no_margins(wordDoc, p, style='Normal'):
        p.style = wordDoc.styles[style]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)

    @staticmethod
    def smallcaps_glosses(p, text, lang):
        if lang in PaperParser.rxGlosses:
            text = PaperParser.rxGlosses[lang].sub(lambda m: '$' + m.group(1).lower() + '$', text)
            for run in PaperParser.rxGlossesNonGlosses.findall(text):
                if PaperParser.rxGlosses[lang].search(run) is not None:
                    p.add_run(run).font.small_caps = True
                else:
                    p.add_run(run)
        else:
            p.text = text

    def __init__(self, analyzer):
        self.analyzer = analyzer
        self.templates = {}  # Jinja2 template cache

    def render_jinja_html(self, templateDir, templateFilename, **context):
        """
        Render a flask template without flask context (needed if
        this file is imported from outside the package).
        """
        try:
            template = self.templates[(templateDir, templateFilename)]
        except KeyError:
            template = jinja2.Environment(
                loader=jinja2.FileSystemLoader(templateDir + '/')
            ).get_template(templateFilename)
            self.templates[(templateDir, templateFilename)] = template
        return template.render(context)

    def process_example(self, lang, num, text, trans, wordDoc=None):
        result = self.analyzer.analyze(lang, text)
        if 'IPA' in result:
            result = result['IPA']
        else:
            result = result['default']
        words = []
        glosses = []
        hangingPuncL = ''
        for w in result:
            if len(w) <= 0:
                continue
            wf = w[0]['wf']
            if self.rxPuncR.search(wf) is not None:
                if len(words) > 0:
                    words[-1] += wf
                else:
                    hangingPuncL += wf
                continue
            elif self.rxPuncL.search(wf) is not None:
                hangingPuncL += wf
                continue
            wf = hangingPuncL + wf
            gloss = 'STEM'
            curGlosses = set()
            curWfParts = set()
            curTrans = set()
            for ana in w:
                curWfParts.add(ana['wfGlossed'])
                curGlosses.add(ana['gloss'])
                if 'trans_ru' in ana:
                    curTrans.add(ana['trans_ru'])
            curGlosses = [g for g in sorted(curGlosses, key=lambda x: (x.count('-'), len(x), x))
                          if len(g) > 0]
            curWfParts = [p for p in sorted(curWfParts, key=lambda x: (x.count('-'), max(len(p) for p in x.split('-')), x))
                          if len(p) > 0]
            curTrans = [t for t in sorted(curTrans, key=lambda x: (-len(x), x))]
            if len(curGlosses) > 0 and len(curWfParts) > 0:
                wf = hangingPuncL + curWfParts[0]
                gloss = curGlosses[0]
                if len(curTrans) > 0:
                    gloss = gloss.replace('STEM', self.rxStemGloss.sub('.', curTrans[0]))
            hangingPuncL = ''
            if (lang in self.rxEnlitics
                    and self.rxEnlitics[lang].search(wf) is not None
                    and len(words) > 0
                    and len(words[-1]) > 0
                    and self.rxPuncR.search(words[-1][-1]) is None):
                words[-1] += '=' + wf
                glosses[-1] += '=' + gloss
            else:
                words.append(wf)
                glosses.append(gloss)

        if wordDoc is not None:
            nCharsWords = len(''.join(w.strip() for w in words))
            nCharsGloss = len(''.join(re.sub('[а-яё][а-яё ,.\\-()]+',
                                             'XXXXXX', g.strip()) for g in glosses))
            nRows = max(nCharsWords // 56, nCharsGloss // 76) + 1
            nCols = 1 + math.ceil(len(words) / nRows)
            table = wordDoc.add_table(rows=nRows * 2 + 1, cols=nCols)
            p = table.cell(0, 0).paragraphs[0]
            p.text = '(' + str(num) + ')'
            PaperParser.p_no_margins(wordDoc, p)
            for iRow in range(nRows):
                p = table.cell(iRow + 1, 0).paragraphs[0]
                PaperParser.p_no_margins(wordDoc, p)
                p = table.cell(iRow + 2, 0).paragraphs[0]
                PaperParser.p_no_margins(wordDoc, p)
            for iCell in range(len(words)):
                if iCell >= len(glosses):
                    break
                iRow = iCell // (nCols - 1)
                iCol = iCell - iRow * (nCols - 1) + 1
                if iCell >= 1:
                    table.cell(nRows * 2, 1).merge(table.cell(nRows * 2, iCol))
                p = table.cell(iRow * 2, iCol).paragraphs[0]
                p.add_run(words[iCell].strip()).italic = True
                PaperParser.p_no_margins(wordDoc, p)
                p = table.cell(iRow * 2 + 1, iCol).paragraphs[0]
                p.style = wordDoc.styles['Gloss']
                PaperParser.p_no_margins(wordDoc, p, 'Gloss')
                if re.search('^(?:[ /*?!.,()_-]*|\\[S[0-9]+\\]:?)$', words[iCell].strip()) is not None:
                    continue
                PaperParser.smallcaps_glosses(p, glosses[iCell].strip(), lang)
            p = table.cell(nRows * 2, 1).paragraphs[0]
            p.text = trans
            PaperParser.p_no_margins(wordDoc, p)
            # self.set_cell_margins(table, 0, 0)
            table.autofit = True

        return self.render_jinja_html('web_app/templates',
                                      'analysis_paper.html',
                                      num=num,
                                      words=words,
                                      glosses=glosses,
                                      translation=trans).strip()

    def analyze(self, lang, text):
        if lang not in self.analyzer.langs:
            return text
        text = '\n' + text.strip() + '\n'
        text = PaperParser.clean_punc(text)
        segments = self.rxExamples.findall(text)
        textProcessed = ''
        wordDoc = Document()
        glossStyle = wordDoc.styles.add_style('Gloss', WD_STYLE_TYPE.PARAGRAPH)
        headerStyle = wordDoc.styles.add_style('Section header', WD_STYLE_TYPE.PARAGRAPH)
        normalStyle = wordDoc.styles['Normal']
        normalStyle.font.name = 'Brill'
        normalStyle.font.size = Pt(10)
        normalStyle.paragraph_format.space_after = Cm(0)
        glossStyle.font.name = 'Brill'
        glossStyle.font.size = Pt(9)
        headerStyle.font.name = 'Brill'
        headerStyle.font.size = Pt(10)
        prevTitle = True
        prevExample = False
        for seg in segments:
            if len(seg[1]) == 0 and len(seg[3]) == 0:
                textProcessed += '<br>'
            elif len(seg[3]) > 0 and len(seg[1]) <= 0:
                para = seg[3]
                if lang in self.rxWordLang:
                    if 'IPA' in self.analyzer.langs[lang]['translit']:
                        transliterator = self.analyzer.langs[lang]['translit']['IPA']
                    else:
                        transliterator = lambda s: s
                    para = self.rxWordLang[lang].sub(lambda m: '<i>' + transliterator(m.group(0)) + '</i>', para)
                textProcessed += '<p>' + para.replace('\n', '</p>\n<p>')[:-3]
                paraRuns = re.findall('<i>.+?</i>|(?:[^<]|<[^i])+', para.strip('\r\n'))
                if len(paraRuns) > 1 or (len(paraRuns) == 1
                                         and re.search('^(?:[ \r\n]*|<i> *</i>[ \r\n]*)$',
                                                       paraRuns[0], flags=re.DOTALL) is None):
                    p = wordDoc.add_paragraph('')
                    p.style = wordDoc.styles['Normal']
                    prevExample = False
                    if not prevTitle:
                        p.paragraph_format.first_line_indent = Cm(1)
                    p.paragraph_format.space_before = Cm(0)
                    p.paragraph_format.space_after = Cm(0)
                    if (len(paraRuns) == 1
                            and re.search('^[^<>]{0,65}[^.?!:;)<> -] *$', paraRuns[0]) is not None
                            and not paraRuns[0].startswith('Table')):
                        p.add_run('XX.X ' + paraRuns[0]).bold = True
                        p.style = wordDoc.styles['Section header']
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.first_line_indent = Cm(0)
                        prevTitle = True
                    else:
                        prevTitle = False
                        for paraRun in paraRuns:
                            if paraRun.startswith('<i>'):
                                p.add_run(paraRun[3:len(paraRun)-4]).italic = True
                            else:
                                PaperParser.smallcaps_glosses(p, paraRun, lang)
            else:
                print(seg)
                if not prevExample:
                    p = wordDoc.add_paragraph('')
                    PaperParser.p_no_margins(wordDoc, p)
                prevExample = True
                textProcessed += self.process_example(lang, seg[0], seg[1], seg[2], wordDoc)
                p = wordDoc.add_paragraph('')
                PaperParser.p_no_margins(wordDoc, p)
        if not os.path.exists('docx'):
            os.makedirs('docx')
        wordDoc.save('docx/processed.docx')
        return textProcessed

